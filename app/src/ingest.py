"""Universal ingestion pipelines for E2-6.

Extractors turn supported binary/text formats into plain text that can be
chunked and fed to the router.  Heavy engines (OCR, speech-to-text) are
optional; if they are missing and the workspace requires local-only processing,
the pipeline returns an honest error instead of leaking content to a cloud model.
"""

from __future__ import annotations

import base64
import io
import json
import sqlite3
from typing import Any

from .context import Context
from .db import get_object, get_routing_policy
from .storage import get_object_store


class IngestError(Exception):
    """Raised when an object cannot be ingested."""


class LocalOnlyEngineMissingError(IngestError):
    """Raised when a local engine is required but not installed."""


_INGEST_MIME_MAP = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/json": "json",
    "application/sql": "sql",
    "text/x-sql": "sql",
    "application/x-sql": "sql",
    "text/csv": "csv",
    "application/csv": "csv",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "text/markdown": "text",
    "text/plain": "text",
    "text/html": "text",
    "application/x-msaccess": "access",
    "application/msaccess": "access",
    "image/png": "image",
    "image/jpeg": "image",
    "image/gif": "image",
    "image/webp": "image",
    "image/bmp": "image",
    "image/tiff": "image",
    "audio/mpeg": "audio",
    "audio/mp4": "audio",
    "audio/wav": "audio",
    "audio/x-wav": "audio",
    "audio/ogg": "audio",
    "audio/webm": "audio",
    "video/mp4": "video",
    "video/webm": "video",
    "video/ogg": "video",
    "video/quicktime": "video",
    "video/x-matroska": "video",
}


def _workspace_requires_local_only(ctx: Context, conn: sqlite3.Connection) -> bool:
    policy = get_routing_policy(ctx, conn)
    return bool(policy.get("require_local_only", 0))


def _read_object_blob(ctx: Context, conn: sqlite3.Connection, object_id: str) -> bytes:
    obj = get_object(ctx, conn, object_id)
    if obj is None:
        raise IngestError(f"Object {object_id} not found")
    store = get_object_store()
    try:
        return store.get_object(obj["bucket"], obj["object_key"])
    except Exception as exc:
        raise IngestError(f"Failed to read object {object_id} from storage: {exc}") from exc


def _detect_ingest_type(mime_type: str | None, file_name: str | None) -> str:
    lowered = (mime_type or "").lower()
    if lowered in _INGEST_MIME_MAP:
        return _INGEST_MIME_MAP[lowered]
    # Fallback by extension.
    ext = (file_name or "").rsplit(".", 1)[-1].lower()
    ext_map = {
        "docx": "docx",
        "doc": "docx",
        "json": "json",
        "sql": "sql",
        "csv": "csv",
        "tsv": "csv",
        "xlsx": "xlsx",
        "xls": "xlsx",
        "md": "text",
        "txt": "text",
        "log": "text",
        "html": "text",
        "htm": "text",
        "accdb": "access",
        "mdb": "access",
        "png": "image",
        "jpg": "image",
        "jpeg": "image",
        "gif": "image",
        "webp": "image",
        "bmp": "image",
        "tiff": "image",
        "mp3": "audio",
        "m4a": "audio",
        "wav": "audio",
        "ogg": "audio",
        "weba": "audio",
        "mp4": "video",
        "webm": "video",
        "ogv": "video",
        "mov": "video",
        "mkv": "video",
    }
    return ext_map.get(ext, lowered.split("/")[-1].split("+")[0])


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------


def _extract_docx(blob: bytes) -> str:
    try:
        import docx  # type: ignore[import-untyped]
    except Exception as exc:  # pragma: no cover - optional dependency
        raise LocalOnlyEngineMissingError("python-docx no está instalado") from exc

    document = docx.Document(io.BytesIO(blob))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables_text = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text for cell in row.cells))
        if rows:
            tables_text.append("\n".join(rows))
    parts = paragraphs
    if tables_text:
        parts.append("---TABLAS---")
        parts.extend(tables_text)
    return "\n\n".join(parts)


def _extract_json(blob: bytes) -> str:
    try:
        text = blob.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise IngestError("JSON file is not valid UTF-8") from exc
    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        raise IngestError(f"Invalid JSON: {exc}") from exc
    return json.dumps(data, ensure_ascii=False, indent=2, sort_keys=True)


def _extract_sql(blob: bytes) -> str:
    try:
        return blob.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise IngestError("SQL file is not valid UTF-8") from exc


def _decode_text(blob: bytes, what: str) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return blob.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise IngestError(f"{what} file is not decodable text")


def _extract_csv(blob: bytes) -> str:
    import csv as _csv
    import io as _io

    text = _decode_text(blob, "CSV")
    try:
        sample = text[:4096]
        dialect = _csv.Sniffer().sniff(sample) if sample.strip() else _csv.excel
    except Exception:
        dialect = _csv.excel
    rows = list(_csv.reader(_io.StringIO(text), dialect))
    if not rows:
        return ""
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def _extract_xlsx(blob: bytes) -> str:
    try:
        import openpyxl  # type: ignore[import-untyped]
    except Exception as exc:  # pragma: no cover - optional dependency
        raise LocalOnlyEngineMissingError("openpyxl no está instalado") from exc

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
    except Exception as exc:
        raise IngestError(f"Invalid XLSX: {exc}") from exc

    parts: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell) for cell in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"--- Hoja: {sheet.title} ---")
            parts.extend(rows)
    workbook.close()
    return "\n".join(parts)


def _extract_plain_text(blob: bytes) -> str:
    return _decode_text(blob, "Text")


def _extract_image(blob: bytes, require_local: bool) -> str:
    try:
        from PIL import Image  # type: ignore[import-untyped]
        import pytesseract  # type: ignore[import-untyped]
    except Exception as exc:  # pragma: no cover - optional dependency
        if require_local:
            raise LocalOnlyEngineMissingError(
                "OCR local no disponible (instalar pytesseract + Pillow)"
            ) from exc
        raise LocalOnlyEngineMissingError("OCR local no disponible") from exc

    try:
        image = Image.open(io.BytesIO(blob))
        return pytesseract.image_to_string(image).strip()
    except Exception as exc:
        raise IngestError(f"OCR failed: {exc}") from exc


def _extract_audio(blob: bytes, require_local: bool) -> str:
    try:
        import whisper  # type: ignore[import-untyped]
    except Exception as exc:  # pragma: no cover - optional dependency
        if require_local:
            raise LocalOnlyEngineMissingError(
                "Transcripción local no disponible (instalar openai-whisper)"
            ) from exc
        raise LocalOnlyEngineMissingError("Transcripción local no disponible") from exc

    try:
        with io.BytesIO(blob) as src:
            # whisper expects a file path or numpy array; passing a file-like works with the API.
            model = whisper.load_model("base")
            result = model.transcribe(src)
        return result.get("text", "").strip()
    except Exception as exc:
        raise IngestError(f"Transcription failed: {exc}") from exc


def _extract_video(blob: bytes, require_local: bool) -> str:
    # Naive first pass: treat as audio extraction via whisper.
    # A production pipeline would separate audio with ffmpeg first.
    return _extract_audio(blob, require_local)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract_text_from_blob(
    *,
    blob: bytes,
    ingest_type: str,
    require_local: bool = False,
) -> str:
    """Return extracted text from a binary blob.

    Raises ``LocalOnlyEngineMissingError`` when a local engine is required but
    not installed; raises ``IngestError`` for malformed or unsupported content.
    """

    if ingest_type == "docx":
        return _extract_docx(blob)
    if ingest_type == "json":
        return _extract_json(blob)
    if ingest_type == "sql":
        return _extract_sql(blob)
    if ingest_type == "csv":
        return _extract_csv(blob)
    if ingest_type == "xlsx":
        return _extract_xlsx(blob)
    if ingest_type == "text":
        return _extract_plain_text(blob)
    if ingest_type == "access":
        raise IngestError(
            "Microsoft Access (.accdb/.mdb) no está soportado: exporta la tabla a CSV o XLSX y súbela de nuevo"
        )
    if ingest_type == "image":
        return _extract_image(blob, require_local)
    if ingest_type == "audio":
        return _extract_audio(blob, require_local)
    if ingest_type == "video":
        return _extract_video(blob, require_local)
    raise IngestError(f"Unsupported ingest type: {ingest_type}")


def extract_text_from_object(
    ctx: Context,
    conn: sqlite3.Connection,
    object_id: str,
) -> dict[str, Any]:
    """Read an object from MinIO and extract ingestible text.

    Returns a dict with ``object_id``, ``file_name``, ``mime_type``,
    ``ingest_type``, ``text`` and ``error`` (None on success).
    """

    obj = get_object(ctx, conn, object_id)
    if obj is None:
        return {"object_id": object_id, "error": "Object not found"}

    blob = _read_object_blob(ctx, conn, object_id)
    ingest_type = _detect_ingest_type(obj.get("mime_type"), obj.get("file_name"))
    require_local = _workspace_requires_local_only(ctx, conn)

    try:
        text = extract_text_from_blob(
            blob=blob,
            ingest_type=ingest_type,
            require_local=require_local,
        )
        return {
            "object_id": object_id,
            "file_name": obj.get("file_name"),
            "mime_type": obj.get("mime_type"),
            "ingest_type": ingest_type,
            "text": text,
            "error": None,
        }
    except LocalOnlyEngineMissingError as exc:
        return {
            "object_id": object_id,
            "file_name": obj.get("file_name"),
            "mime_type": obj.get("mime_type"),
            "ingest_type": ingest_type,
            "text": "",
            "error": f"LOCAL_ENGINE_MISSING: {exc}",
        }
    except IngestError as exc:
        return {
            "object_id": object_id,
            "file_name": obj.get("file_name"),
            "mime_type": obj.get("mime_type"),
            "ingest_type": ingest_type,
            "text": "",
            "error": f"INGEST_ERROR: {exc}",
        }


# Máximo para adjuntos de imagen enviados a modelos con visión.
# Anthropic acepta hasta ~5 MB por imagen; OpenAI hasta 20 MB en data URL.
IMAGE_ATTACHMENT_MAX_BYTES = 5 * 1024 * 1024

_VISION_MIME_NORMALIZE = {"image/jpg": "image/jpeg"}
_VISION_SUPPORTED_MIMES = {"image/png", "image/jpeg", "image/gif", "image/webp"}


def load_image_attachment(
    ctx: Context,
    conn: sqlite3.Connection,
    object_id: str,
) -> dict[str, Any]:
    """Read an image object from storage and return it base64-encoded.

    Returns a dict with ``object_id``, ``file_name``, ``mime_type``,
    ``data_b64`` and ``error`` (None on success). Used to build multimodal
    completion messages for vision-capable models.
    """

    obj = get_object(ctx, conn, object_id)
    if obj is None:
        return {"object_id": object_id, "error": "Object not found"}

    file_name = obj.get("file_name")
    mime = (obj.get("mime_type") or "").lower()
    mime = _VISION_MIME_NORMALIZE.get(mime, mime)
    if mime not in _VISION_SUPPORTED_MIMES:
        return {
            "object_id": object_id,
            "file_name": file_name,
            "mime_type": mime,
            "error": f"Unsupported image type for vision: {mime or 'unknown'}",
        }

    try:
        blob = _read_object_blob(ctx, conn, object_id)
    except IngestError as exc:
        return {
            "object_id": object_id,
            "file_name": file_name,
            "mime_type": mime,
            "error": f"INGEST_ERROR: {exc}",
        }

    if len(blob) > IMAGE_ATTACHMENT_MAX_BYTES:
        return {
            "object_id": object_id,
            "file_name": file_name,
            "mime_type": mime,
            "error": (
                f"Image too large ({len(blob)} bytes; max {IMAGE_ATTACHMENT_MAX_BYTES}). "
                "Redimensiona la imagen e inténtalo de nuevo."
            ),
        }

    return {
        "object_id": object_id,
        "file_name": file_name,
        "mime_type": mime,
        "data_b64": base64.b64encode(blob).decode("ascii"),
        "error": None,
    }
